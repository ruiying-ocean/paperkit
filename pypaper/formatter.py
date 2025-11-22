"""
Formatting module for Word documents
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .config import DEFAULT_CONFIG, PAPER_SIZES


def apply_formatting(input_file, output_file=None, config=None):
    """
    Apply formatting to an existing Word document.

    Args:
        input_file: Path to input .docx file
        output_file: Path to output .docx file (if None, overwrites input)
        config: Optional dict with custom settings

    Returns:
        True if successful, False otherwise
    """

    cfg = DEFAULT_CONFIG.copy()
    if config:
        cfg.update(config)

    if output_file is None:
        output_file = input_file

    print(f"Loading: {input_file}")
    doc = Document(input_file)

    # Set page size and margins
    paper_size = cfg.get('paper_size', 'a4')
    print(f"Setting page size to {paper_size.upper()} and margins to {cfg['margins']} inch...")
    for section in doc.sections:
        # Get paper size from config
        if paper_size in PAPER_SIZES:
            width, height = PAPER_SIZES[paper_size]
            section.page_width = Inches(width)
            section.page_height = Inches(height)
        else:
            # Default to A4 if invalid
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)

        # Set margins
        margin = Inches(cfg['margins'])
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

    # Apply formatting to all paragraphs
    print(f"Applying font: {cfg['font']}, {cfg['font_size']}pt...")
    print(f"Setting line spacing: {cfg['line_spacing']}...")

    paragraph_count = 0
    for paragraph in doc.paragraphs:
        paragraph_count += 1

        # Apply font to all runs
        for run in paragraph.runs:
            run.font.name = cfg['font']
            run.font.size = Pt(cfg['font_size'])
            run.font.color.rgb = RGBColor(0, 0, 0)  # Explicit black

        # Apply heading styles
        style_name = paragraph.style.name

        if style_name == 'Title' or 'Title' in style_name:
            for run in paragraph.runs:
                run.font.size = Pt(cfg['title_size'])
                run.font.bold = True
                run.font.name = cfg['font']
                run.font.color.rgb = RGBColor(0, 0, 0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif style_name.startswith('Heading 1'):
            for run in paragraph.runs:
                run.font.size = Pt(cfg['heading1_size'])
                run.font.bold = True
                run.font.name = cfg['font']
                run.font.color.rgb = RGBColor(0, 0, 0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif style_name.startswith('Heading 2'):
            for run in paragraph.runs:
                run.font.size = Pt(cfg['heading2_size'])
                run.font.bold = True
                run.font.name = cfg['font']
                run.font.color.rgb = RGBColor(0, 0, 0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif style_name.startswith('Heading 3'):
            for run in paragraph.runs:
                run.font.size = Pt(cfg['heading3_size'])
                run.font.bold = True
                run.font.name = cfg['font']
                run.font.color.rgb = RGBColor(0, 0, 0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Set line spacing
        paragraph.paragraph_format.line_spacing = cfg['line_spacing']
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(6)

    # Format tables
    table_count = 0
    for table in doc.tables:
        table_count += 1
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = cfg['font']
                        run.font.size = Pt(max(9, cfg['font_size'] - 2))
                        run.font.color.rgb = RGBColor(0, 0, 0)

    # Save document
    print(f"Saving: {output_file}")
    doc.save(output_file)

    print()
    print("=" * 60)
    print("✓ Formatting applied successfully!")
    print("=" * 60)
    print(f"Processed: {paragraph_count} paragraphs, {table_count} tables")
    print(f"Font:      {cfg['font']}, {cfg['font_size']}pt")
    print(f"Title:     {cfg['title_size']}pt, Bold, Black")
    print(f"Heading 1: {cfg['heading1_size']}pt, Bold, Black")
    print(f"Heading 2: {cfg['heading2_size']}pt, Bold, Black")
    print(f"Spacing:   {cfg['line_spacing']}")
    print(f"Margins:   {cfg['margins']} inch")
    print()

    return True
