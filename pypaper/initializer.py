"""
Paper initialization module
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from .config import DEFAULT_CONFIG, PAPER_SIZES
from .templates import get_template


def init_paper(title, output_file="paper.docx", config=None, template=None):
    """
    Create a new formatted academic paper.

    Args:
        title: Paper title
        output_file: Path to save the document
        config: Optional formatting configuration
        template: Journal template name (e.g., 'agu', 'nature', 'science')

    Returns:
        True if successful, False otherwise
    """

    # Start with default config
    cfg = DEFAULT_CONFIG.copy()

    # Apply journal template if specified
    if template:
        template_config = get_template(template)
        cfg.update(template_config)

    # Apply custom config overrides
    if config:
        cfg.update(config)

    print()
    print("=" * 60)
    print("Initializing New Academic Paper")
    print("=" * 60)
    print(f"Title: {title}")
    print(f"File:  {output_file}")
    print()

    # Create document
    doc = Document()

    # Set page size and margins
    for section in doc.sections:
        # Get paper size from config
        paper_size = cfg.get('paper_size', 'a4')
        if paper_size in PAPER_SIZES:
            width, height = PAPER_SIZES[paper_size]
            section.page_width = Inches(width)
            section.page_height = Inches(height)
        else:
            # Default to A4 if invalid
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)

        # Set margins
        margin = Inches(cfg['margins'])
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

    # Add Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.name = cfg['font']
    title_run.font.size = Pt(cfg['title_size'])
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_run.font.underline = False  # Explicitly no underline
    title_para.paragraph_format.line_spacing = cfg['line_spacing']

    # Add Authors (placeholder)
    author_para = doc.add_paragraph()
    author_run = author_para.add_run("Author Name¹, Second Author², Third Author¹")
    author_run.font.name = cfg['font']
    author_run.font.size = Pt(cfg['font_size'])
    author_run.font.color.rgb = RGBColor(0, 0, 0)
    author_para.paragraph_format.line_spacing = cfg['line_spacing']

    # Add affiliations
    doc.add_paragraph()
    affil_para = doc.add_paragraph()
    affil_run = affil_para.add_run(
        "¹ School of Environmental Sciences, University of East Anglia, Norwich, UK\n"
        "² Department, Institution, City, Country"
    )
    affil_run.font.name = cfg['font']
    affil_run.font.size = Pt(cfg['font_size'] - 1)
    affil_run.font.color.rgb = RGBColor(0, 0, 0)
    affil_para.paragraph_format.line_spacing = cfg['line_spacing']

    doc.add_paragraph()

    # Get sections from template or use default
    if 'sections' in cfg:
        section_titles = cfg['sections']
    else:
        section_titles = [
            'Abstract', 'Introduction', 'Methods', 'Results',
            'Discussion', 'Conclusions', 'Acknowledgements',
            'Data Availability', 'Author Contributions',
            'Competing Interests', 'References'
        ]

    # Create section list with placeholders
    sections = []
    for section_title in section_titles:
        if section_title.lower() == 'abstract':
            placeholder = '[Write your abstract here. Typically 150-250 words.]'
        elif 'introduction' in section_title.lower():
            placeholder = '[Introduce the research question and background.]'
        elif 'method' in section_title.lower():
            placeholder = '[Describe your methodology.]'
        elif 'result' in section_title.lower():
            placeholder = '[Present your findings.]'
        elif 'discussion' in section_title.lower():
            placeholder = '[Interpret results and discuss implications.]'
        elif 'conclusion' in section_title.lower():
            placeholder = '[Summarize main findings.]'
        elif 'reference' in section_title.lower():
            placeholder = '[References will be added here.]'
        elif 'competing' in section_title.lower() or 'conflict' in section_title.lower():
            placeholder = 'The authors declare no competing interests.'
        else:
            placeholder = f'[Content for {section_title} section.]'
        sections.append((section_title, placeholder))

    for section_title, placeholder_text in sections:
        # Add heading
        heading = doc.add_heading(section_title, level=1)
        for run in heading.runs:
            run.font.name = cfg['font']
            run.font.size = Pt(cfg['heading1_size'])
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Add placeholder
        para = doc.add_paragraph(placeholder_text)
        para.paragraph_format.line_spacing = cfg['line_spacing']
        for run in para.runs:
            run.font.name = cfg['font']
            run.font.size = Pt(cfg['font_size'])
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black

    # Save
    doc.save(output_file)

    print("✓ Paper initialized successfully!")
    print()
    print("Sections created:")
    for section_title, _ in sections:
        print(f"  • {section_title}")
    print()
    print(f"Font:     {cfg['font']}, {cfg['font_size']}pt")
    print(f"Spacing:  {cfg['line_spacing']}")
    print(f"Margins:  {cfg['margins']} inch")
    print()

    return True
