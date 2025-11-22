"""
Formatting module for Word documents
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .config import DEFAULT_CONFIG, PAPER_SIZES


def clear_table_borders(table):
    """
    Clear all table-level borders.

    Args:
        table: python-docx table object
    """
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove table borders element if it exists
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)

    # Create new empty borders element
    tblBorders = OxmlElement('w:tblBorders')
    tblPr.append(tblBorders)

    # Set all table-level borders to none
    no_border = {'sz': '0', 'val': 'none', 'color': 'auto'}
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        edge_element = OxmlElement(f'w:{edge}')
        for key, value in no_border.items():
            edge_element.set(qn(f'w:{key}'), value)
        tblBorders.append(edge_element)


def set_cell_border(cell, **kwargs):
    """
    Set cell border for APA table style.

    Args:
        cell: Table cell object
        **kwargs: Border positions (top, bottom, left, right, insideH, insideV)
                  with values like {'sz': 12, 'val': 'single', 'color': '000000'}
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Get or create table cell borders element
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # Set borders for specified positions
    for edge, border_props in kwargs.items():
        edge_element = tcBorders.find(qn(f'w:{edge}'))
        if edge_element is None:
            edge_element = OxmlElement(f'w:{edge}')
            tcBorders.append(edge_element)

        # Set border properties
        for key, value in border_props.items():
            edge_element.set(qn(f'w:{key}'), str(value))


def apply_apa_table_style(table, cfg):
    """
    Apply APA 7th edition table formatting (3-line table).

    APA 3-line table rules:
    - Line 1: Top of table (above header row)
    - Line 2: Bottom of header row
    - Line 3: Bottom of table
    - No other lines (no vertical lines, no lines between data rows)

    Args:
        table: python-docx table object
        cfg: Configuration dictionary
    """
    # First, clear all table-level borders
    clear_table_borders(table)

    # APA border style: single line, black, 1/2 pt width
    border_style = {'sz': 6, 'val': 'single', 'color': '000000'}
    no_border = {'sz': 0, 'val': 'none'}

    num_rows = len(table.rows)

    for row_idx, row in enumerate(table.rows):
        is_first_row = (row_idx == 0)
        is_last_row = (row_idx == num_rows - 1)

        for cell in row.cells:
            # Format cell text
            for paragraph in cell.paragraphs:
                # Left align
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                for run in paragraph.runs:
                    run.font.name = cfg['font']
                    run.font.size = Pt(cfg['font_size'])
                    run.font.color.rgb = RGBColor(0, 0, 0)

                    # Bold header row
                    if is_first_row:
                        run.font.bold = True

            # First, clear ALL borders for this cell
            set_cell_border(cell,
                top=no_border,
                bottom=no_border,
                left=no_border,
                right=no_border,
                insideH=no_border,
                insideV=no_border
            )

            # Then apply only the APA 3-line borders
            # Top border: Line 1 (only on first row)
            if is_first_row:
                set_cell_border(cell, top=border_style)

            # Bottom border: Line 2 (bottom of header) and Line 3 (bottom of table)
            if is_first_row or is_last_row:
                set_cell_border(cell, bottom=border_style)


def format_table_caption(paragraph, table_number, cfg):
    """
    Format a paragraph as an APA-style table caption.

    APA caption format:
    - "Table X" in bold, italics
    - Title in italics only
    - Left-aligned
    - Above the table

    Args:
        paragraph: Paragraph object containing caption text
        table_number: Integer table number
        cfg: Configuration dictionary
    """
    text = paragraph.text.strip()

    # Clear existing runs
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)

    # Check if text already starts with "Table X"
    if not text.lower().startswith('table'):
        # Add "Table X" prefix
        table_label = paragraph.add_run(f"Table {table_number}")
        table_label.font.bold = True
        table_label.font.italic = True
        table_label.font.name = cfg['font']
        table_label.font.size = Pt(cfg['font_size'])
        table_label.font.color.rgb = RGBColor(0, 0, 0)  # Black text

        # Add line break
        paragraph.add_run("\n")

        # Add title text
        title_run = paragraph.add_run(text)
        title_run.font.italic = True
        title_run.font.name = cfg['font']
        title_run.font.size = Pt(cfg['font_size'])
        title_run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
    else:
        # Text already has "Table X", just format it
        title_run = paragraph.add_run(text)
        title_run.font.italic = True
        title_run.font.name = cfg['font']
        title_run.font.size = Pt(cfg['font_size'])
        title_run.font.color.rgb = RGBColor(0, 0, 0)  # Black text

    # Left align
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(0)


def apply_formatting(input_file, output_file=None, config=None):
    """
    Apply formatting to an existing Word document.

    Args:
        input_file: Path to input .docx file
        output_file: Path to output .docx file (if None, overwrites input)
        config: Optional dict with custom settings

    Returns:
        True if successful, False otherwise
    """

    cfg = DEFAULT_CONFIG.copy()
    if config:
        cfg.update(config)

    if output_file is None:
        output_file = input_file

    print(f"Loading: {input_file}")
    doc = Document(input_file)

    # Set page size and margins
    paper_size = cfg.get('paper_size', 'a4')
    print(f"Setting page size to {paper_size.upper()} and margins to {cfg['margins']} inch...")
    for section in doc.sections:
        # Get paper size from config
        if paper_size in PAPER_SIZES:
            width, height = PAPER_SIZES[paper_size]
            section.page_width = Inches(width)
            section.page_height = Inches(height)
        else:
            # Default to A4 if invalid
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)

        # Set margins
        margin = Inches(cfg['margins'])
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

    # Apply formatting to all paragraphs
    print(f"Applying font: {cfg['font']}, {cfg['font_size']}pt...")
    print(f"Setting line spacing: {cfg['line_spacing']}...")

    paragraph_count = 0
    for paragraph in doc.paragraphs:
        paragraph_count += 1

        # Apply font to all runs
        for run in paragraph.runs:
            run.font.name = cfg['font']
            run.font.size = Pt(cfg['font_size'])
            run.font.color.rgb = RGBColor(0, 0, 0)  # Explicit black

        # Apply heading styles
        style_name = paragraph.style.name

        if style_name == 'Title' or 'Title' in style_name:
            for run in paragraph.runs:
                run.font.size = Pt(cfg['title_size'])
                run.font.bold = True
                run.font.name = cfg['font']
                run.font.color.rgb = RGBColor(0, 0, 0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif style_name.startswith('Heading 1'):
            for run in paragraph.runs:
                run.font.size = Pt(cfg['heading1_size'])
                run.font.bold = True
                run.font.name = cfg['font']
                run.font.color.rgb = RGBColor(0, 0, 0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif style_name.startswith('Heading 2'):
            for run in paragraph.runs:
                run.font.size = Pt(cfg['heading2_size'])
                run.font.bold = True
                run.font.name = cfg['font']
                run.font.color.rgb = RGBColor(0, 0, 0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif style_name.startswith('Heading 3'):
            for run in paragraph.runs:
                run.font.size = Pt(cfg['heading3_size'])
                run.font.bold = True
                run.font.name = cfg['font']
                run.font.color.rgb = RGBColor(0, 0, 0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Set line spacing
        paragraph.paragraph_format.line_spacing = cfg['line_spacing']
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(6)

    # Format tables with APA style
    print("Applying APA table formatting...")
    table_count = 0
    table_number = 1

    # Process tables and their captions
    # Strategy: Look for paragraphs with "Table" or table captions before each table
    for i, element in enumerate(doc.element.body):
        # Check if this is a table element
        if element.tag.endswith('tbl'):
            table_count += 1

            # Get the table object
            for table in doc.tables:
                if table._element == element:
                    # Look for caption in previous paragraph
                    if i > 0:
                        prev_element = doc.element.body[i - 1]
                        if prev_element.tag.endswith('p'):
                            # Find the corresponding paragraph object
                            for para in doc.paragraphs:
                                if para._element == prev_element:
                                    text = para.text.strip().lower()
                                    # Check if this looks like a table caption
                                    if (text.startswith('table') or
                                        para.style.name == 'Caption' or
                                        'caption' in para.style.name.lower()):
                                        format_table_caption(para, table_number, cfg)

                    # Apply APA table style
                    apply_apa_table_style(table, cfg)
                    table_number += 1
                    break

    # Save document
    print(f"Saving: {output_file}")
    doc.save(output_file)

    print()
    print("=" * 60)
    print("✓ Formatting applied successfully!")
    print("=" * 60)
    print(f"Processed: {paragraph_count} paragraphs, {table_count} tables")
    print(f"Font:      {cfg['font']}, {cfg['font_size']}pt")
    print(f"Title:     {cfg['title_size']}pt, Bold, Black")
    print(f"Heading 1: {cfg['heading1_size']}pt, Bold, Black")
    print(f"Heading 2: {cfg['heading2_size']}pt, Bold, Black")
    print(f"Spacing:   {cfg['line_spacing']}")
    print(f"Margins:   {cfg['margins']} inch")
    print()

    return True
