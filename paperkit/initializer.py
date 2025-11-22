"""
Paper initialization module
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from .config import DEFAULT_CONFIG, PAPER_SIZES
from .templates import get_template
from .formatter import apply_apa_table_style


def init_paper(title, output_file="paper.docx", config=None, template=None):
    """
    Create a new formatted academic paper.

    Args:
        title: Paper title
        output_file: Path to save the document
        config: Optional formatting configuration
        template: Journal template name (e.g., 'agu', 'nature', 'science')

    Returns:
        True if successful, False otherwise
    """

    # Start with default config
    cfg = DEFAULT_CONFIG.copy()

    # Apply journal template if specified
    if template:
        template_config = get_template(template)
        cfg.update(template_config)

    # Apply custom config overrides
    if config:
        cfg.update(config)

    print()
    print("=" * 60)
    print("Initializing New Academic Paper")
    print("=" * 60)
    print(f"Title: {title}")
    print(f"File:  {output_file}")
    print()

    # Create document
    doc = Document()

    # Set page size and margins
    for section in doc.sections:
        # Get paper size from config
        paper_size = cfg.get('paper_size', 'a4')
        if paper_size in PAPER_SIZES:
            width, height = PAPER_SIZES[paper_size]
            section.page_width = Inches(width)
            section.page_height = Inches(height)
        else:
            # Default to A4 if invalid
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)

        # Set margins
        margin = Inches(cfg['margins'])
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

    # Add Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.name = cfg['font']
    title_run.font.size = Pt(cfg['title_size'])
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_run.font.underline = False  # Explicitly no underline
    title_para.paragraph_format.line_spacing = cfg['line_spacing']

    # Add Authors (placeholder)
    author_para = doc.add_paragraph()
    author_run = author_para.add_run("Author Name¹, Second Author², Third Author¹")
    author_run.font.name = cfg['font']
    author_run.font.size = Pt(cfg['font_size'])
    author_run.font.color.rgb = RGBColor(0, 0, 0)
    author_para.paragraph_format.line_spacing = cfg['line_spacing']

    # Add affiliations
    doc.add_paragraph()
    affil_para = doc.add_paragraph()
    affil_run = affil_para.add_run(
        "¹ School of Environmental Sciences, University of East Anglia, Norwich, UK\n"
        "² Department, Institution, City, Country"
    )
    affil_run.font.name = cfg['font']
    affil_run.font.size = Pt(cfg['font_size'] - 1)
    affil_run.font.color.rgb = RGBColor(0, 0, 0)
    affil_para.paragraph_format.line_spacing = cfg['line_spacing']

    doc.add_paragraph()

    # Get sections from template or use default
    if 'sections' in cfg:
        section_titles = cfg['sections']
    else:
        section_titles = [
            'Abstract', 'Introduction', 'Methods', 'Results',
            'Discussion', 'Conclusions', 'Acknowledgements',
            'Data Availability', 'Author Contributions',
            'Competing Interests', 'References'
        ]

    # Create section list with placeholders
    sections = []
    for section_title in section_titles:
        if section_title.lower() == 'abstract':
            placeholder = '[Write your abstract here. Typically 150-250 words.]'
        elif 'introduction' in section_title.lower():
            placeholder = '[Introduce the research question and background.]'
        elif 'method' in section_title.lower():
            placeholder = '[Describe your methodology.]'
        elif 'result' in section_title.lower():
            placeholder = '[Present your findings.]'
        elif 'discussion' in section_title.lower():
            placeholder = '[Interpret results and discuss implications.]'
        elif 'conclusion' in section_title.lower():
            placeholder = '[Summarize main findings.]'
        elif 'reference' in section_title.lower():
            placeholder = '[References will be added here.]'
        elif 'competing' in section_title.lower() or 'conflict' in section_title.lower():
            placeholder = 'The authors declare no competing interests.'
        else:
            placeholder = f'[Content for {section_title} section.]'
        sections.append((section_title, placeholder))

    for section_title, placeholder_text in sections:
        # Add heading
        heading = doc.add_heading(section_title, level=1)
        for run in heading.runs:
            run.font.name = cfg['font']
            run.font.size = Pt(cfg['heading1_size'])
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Add sample text with chemical formulas in Introduction section
        if 'introduction' in section_title.lower():
            chem_para = doc.add_paragraph()
            chem_para.paragraph_format.line_spacing = cfg['line_spacing']

            # Add text with chemical formulas
            run1 = chem_para.add_run('Ocean acidification due to increased atmospheric CO')
            run1.font.name = cfg['font']
            run1.font.size = Pt(cfg['font_size'])
            run1.font.color.rgb = RGBColor(0, 0, 0)

            # CO2 subscript
            run2 = chem_para.add_run('2')
            run2.font.name = cfg['font']
            run2.font.size = Pt(cfg['font_size'])
            run2.font.subscript = True
            run2.font.color.rgb = RGBColor(0, 0, 0)

            run3 = chem_para.add_run(' concentrations affects calcium carbonate (CaCO')
            run3.font.name = cfg['font']
            run3.font.size = Pt(cfg['font_size'])
            run3.font.color.rgb = RGBColor(0, 0, 0)

            # CaCO3 subscript
            run4 = chem_para.add_run('3')
            run4.font.name = cfg['font']
            run4.font.size = Pt(cfg['font_size'])
            run4.font.subscript = True
            run4.font.color.rgb = RGBColor(0, 0, 0)

            run5 = chem_para.add_run(') saturation states. The pH of surface waters has decreased, with measurements showing changes in HCO')
            run5.font.name = cfg['font']
            run5.font.size = Pt(cfg['font_size'])
            run5.font.color.rgb = RGBColor(0, 0, 0)

            # HCO3- subscript and superscript
            run6 = chem_para.add_run('3')
            run6.font.name = cfg['font']
            run6.font.size = Pt(cfg['font_size'])
            run6.font.subscript = True
            run6.font.color.rgb = RGBColor(0, 0, 0)

            run7 = chem_para.add_run('−')
            run7.font.name = cfg['font']
            run7.font.size = Pt(cfg['font_size'])
            run7.font.superscript = True
            run7.font.color.rgb = RGBColor(0, 0, 0)

            run8 = chem_para.add_run(' and CO')
            run8.font.name = cfg['font']
            run8.font.size = Pt(cfg['font_size'])
            run8.font.color.rgb = RGBColor(0, 0, 0)

            run9 = chem_para.add_run('3')
            run9.font.name = cfg['font']
            run9.font.size = Pt(cfg['font_size'])
            run9.font.subscript = True
            run9.font.color.rgb = RGBColor(0, 0, 0)

            run10 = chem_para.add_run('2−')
            run10.font.name = cfg['font']
            run10.font.size = Pt(cfg['font_size'])
            run10.font.superscript = True
            run10.font.color.rgb = RGBColor(0, 0, 0)

            run11 = chem_para.add_run(' concentrations across different ocean regions.')
            run11.font.name = cfg['font']
            run11.font.size = Pt(cfg['font_size'])
            run11.font.color.rgb = RGBColor(0, 0, 0)

        # Add subheadings for Methods section
        elif 'method' in section_title.lower():
            # Subheading: Study Area
            subheading1 = doc.add_heading('Study Area', level=2)
            for run in subheading1.runs:
                run.font.name = cfg['font']
                run.font.size = Pt(cfg['heading2_size'])
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            sub_para1 = doc.add_paragraph('[Describe the study area and location.]')
            sub_para1.paragraph_format.line_spacing = cfg['line_spacing']
            for run in sub_para1.runs:
                run.font.name = cfg['font']
                run.font.size = Pt(cfg['font_size'])
                run.font.color.rgb = RGBColor(0, 0, 0)

            doc.add_paragraph()

            # Subheading: Data Collection
            subheading2 = doc.add_heading('Data Collection', level=2)
            for run in subheading2.runs:
                run.font.name = cfg['font']
                run.font.size = Pt(cfg['heading2_size'])
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            sub_para2 = doc.add_paragraph('[Describe data collection procedures.]')
            sub_para2.paragraph_format.line_spacing = cfg['line_spacing']
            for run in sub_para2.runs:
                run.font.name = cfg['font']
                run.font.size = Pt(cfg['font_size'])
                run.font.color.rgb = RGBColor(0, 0, 0)

        # Add sample table in Results section
        elif 'result' in section_title.lower():
            # Subheading: Descriptive Statistics
            subheading_stats = doc.add_heading('Descriptive Statistics', level=2)
            for run in subheading_stats.runs:
                run.font.name = cfg['font']
                run.font.size = Pt(cfg['heading2_size'])
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            # Add table caption
            caption = doc.add_paragraph('Sample Descriptive Statistics')
            caption_run = caption.runs[0]
            caption_run.font.name = cfg['font']
            caption_run.font.size = Pt(cfg['font_size'])
            caption_run.font.bold = True
            caption_run.font.italic = True
            caption_run.font.color.rgb = RGBColor(0, 0, 0)
            caption_run.text = 'Table 1\nSample Descriptive Statistics'
            caption.paragraph_format.space_before = Pt(12)
            caption.paragraph_format.space_after = Pt(0)

            # Create sample table
            table = doc.add_table(rows=4, cols=4)

            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Variable'
            header_cells[1].text = 'M'
            header_cells[2].text = 'SD'
            header_cells[3].text = 'N'

            # Data rows
            data_rows = [
                ('Age', '35.2', '8.4', '120'),
                ('Experience (years)', '10.5', '4.2', '120'),
                ('Performance Score', '78.3', '12.1', '120')
            ]

            for i, (var, mean, sd, n) in enumerate(data_rows, 1):
                row = table.rows[i]
                row.cells[0].text = var
                row.cells[1].text = mean
                row.cells[2].text = sd
                row.cells[3].text = n

            # Apply APA table formatting
            apply_apa_table_style(table, cfg)

            # Add note below table
            doc.add_paragraph()
            note = doc.add_paragraph('Note. M = Mean; SD = Standard Deviation; N = Sample size.')
            note_run = note.runs[0]
            note_run.font.name = cfg['font']
            note_run.font.size = Pt(cfg['font_size'])
            note_run.font.italic = True
            note_run.font.color.rgb = RGBColor(0, 0, 0)
            note.paragraph_format.space_before = Pt(0)
            note.paragraph_format.space_after = Pt(6)

            # Add sample figure placeholder
            doc.add_paragraph()
            doc.add_paragraph()

            # Subheading: Spatial Patterns
            subheading_spatial = doc.add_heading('Spatial Patterns', level=2)
            for run in subheading_spatial.runs:
                run.font.name = cfg['font']
                run.font.size = Pt(cfg['heading2_size'])
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            # Add figure image
            fig_para = doc.add_paragraph()
            fig_para.paragraph_format.alignment = 1  # Center alignment

            # Get path to example image
            package_dir = os.path.dirname(os.path.abspath(__file__))
            image_path = os.path.join(package_dir, '..', 'static', 'example.png')

            # Add image if it exists, otherwise use placeholder text
            if os.path.exists(image_path):
                fig_run = fig_para.add_run()
                fig_run.add_picture(image_path, width=Inches(5.0))
            else:
                fig_run = fig_para.add_run('[Insert figure here]')
                fig_run.font.name = cfg['font']
                fig_run.font.size = Pt(cfg['font_size'])
                fig_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray text
                fig_run.font.italic = True

            # Add figure caption with SEQ field for auto-numbering
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            fig_caption = doc.add_paragraph()
            fig_caption.paragraph_format.alignment = 1  # Center alignment
            fig_caption.paragraph_format.space_before = Pt(6)
            fig_caption.paragraph_format.space_after = Pt(12)

            # Add "Figure " text
            caption_run = fig_caption.add_run('Figure ')
            caption_run.font.name = cfg['font']
            caption_run.font.size = Pt(cfg['font_size'])
            caption_run.font.italic = True

            # Add SEQ field for auto-numbering
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'SEQ Figure \\* ARABIC'

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')

            # Add placeholder for the number
            num_run = fig_caption.add_run('1')
            num_run.font.name = cfg['font']
            num_run.font.size = Pt(cfg['font_size'])
            num_run.font.italic = True

            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')

            # Insert field elements into the first run
            r_element = caption_run._element
            r_element.append(fldChar1)
            r_element.append(instrText)
            r_element.append(fldChar2)

            # Insert end field char after the number run
            num_run._element.append(fldChar3)

            # Add caption text
            text_run = fig_caption.add_run('. Sample surface water temperature map')
            text_run.font.name = cfg['font']
            text_run.font.size = Pt(cfg['font_size'])
            text_run.font.italic = True

            # Add example paragraph with cross-reference to the figure
            doc.add_paragraph()
            ref_para = doc.add_paragraph()
            ref_para.paragraph_format.line_spacing = cfg['line_spacing']

            # Add text before the reference
            ref_run1 = ref_para.add_run('The surface water temperature data shows significant variation across regions (see ')
            ref_run1.font.name = cfg['font']
            ref_run1.font.size = Pt(cfg['font_size'])
            ref_run1.font.color.rgb = RGBColor(0, 0, 0)

            # Add cross-reference field to the figure
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'REF _Ref1 \\h'

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')

            # Add placeholder for the reference text
            ref_run2 = ref_para.add_run('Figure 1')
            ref_run2.font.name = cfg['font']
            ref_run2.font.size = Pt(cfg['font_size'])
            ref_run2.font.color.rgb = RGBColor(0, 0, 0)

            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')

            # Insert field elements
            ref_run1._element.append(fldChar1)
            ref_run1._element.append(instrText)
            ref_run1._element.append(fldChar2)
            ref_run2._element.append(fldChar3)

            # Add text after the reference
            ref_run3 = ref_para.add_run('), with temperatures ranging from 0°C to 30°C.')
            ref_run3.font.name = cfg['font']
            ref_run3.font.size = Pt(cfg['font_size'])
            ref_run3.font.color.rgb = RGBColor(0, 0, 0)

            # Add IPCC paragraph with degree symbols
            doc.add_paragraph()
            ipcc_para = doc.add_paragraph()
            ipcc_para.paragraph_format.line_spacing = cfg['line_spacing']

            ipcc_text = 'The likely range of total human-caused global surface temperature increase from 1850–1900 to 2010–2019 is 0.8°C to 1.3°C, with a best estimate of 1.07°C. It is likely that well-mixed GHGs contributed a warming of 1.0°C to 2.0°C, other human drivers (principally aerosols) contributed a cooling of 0.0°C to 0.8°C, natural drivers changed global surface temperature by –0.1°C to +0.1°C, and internal variability changed it by –0.2°C to +0.2°C. It is very likely that well-mixed GHGs were the main driver of tropospheric warming since 1979 and extremely likely that human-caused stratospheric ozone depletion was the main driver of cooling of the lower stratosphere between 1979 and the mid-1990s.'

            ipcc_run = ipcc_para.add_run(ipcc_text)
            ipcc_run.font.name = cfg['font']
            ipcc_run.font.size = Pt(cfg['font_size'])
            ipcc_run.font.color.rgb = RGBColor(0, 0, 0)

        # Add placeholder for all other sections
        else:
            para = doc.add_paragraph(placeholder_text)
            para.paragraph_format.line_spacing = cfg['line_spacing']
            for run in para.runs:
                run.font.name = cfg['font']
                run.font.size = Pt(cfg['font_size'])
                run.font.color.rgb = RGBColor(0, 0, 0)

    # Save
    doc.save(output_file)

    print("✓ Paper initialized successfully!")
    print()
    print("Sections created:")
    for section_title, _ in sections:
        print(f"  • {section_title}")
    print()
    print(f"Font:     {cfg['font']}, {cfg['font_size']}pt")
    print(f"Spacing:  {cfg['line_spacing']}")
    print(f"Margins:  {cfg['margins']} inch")
    print(f"Includes: Sample APA-formatted table and figure in Results section")
    print()

    return True
